import docx

def update_template_var(path):
    print(f"Updating {path}...")
    doc = docx.Document(path)
    
    # We want to replace {{ v.severity }} with {{ v.severity_detail }} ONLY in Table 2 or later (The Detail Table)
    # Assuming Table 0 is Header, Table 1 is Summary, Table 2 is Detail
    
    found_count = 0
    
    if len(doc.tables) > 2:
        for t_idx in range(2, len(doc.tables)):
            table = doc.tables[t_idx]
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{{ v.severity }}" in p.text:
                            # Replace strictly
                            print(f"Found match in T{t_idx}: {p.text}")
                            
                            # Simple string replacement in run text is risky if split across runs
                            # But since user likely typed it or it's a simple placeholder, direct run replacement often works
                            # Trying paragraph text replacement
                            if "{{ v.severity }}" in p.text:
                                p.text = p.text.replace("{{ v.severity }}", "{{ v.severity_detail }}")
                                found_count += 1
                                # Note: modifying p.text usually clears formatting. 
                                # If formatting is critical, we should iterate runs.
                                # But standard templating usually lets the variable carry the formatting (RichText).
                                # Given we are using RichText, the container formatting is less critical.
    
    print(f"Replaced {found_count} occurrences.")
    doc.save(path)
    print("Template updated.")

if __name__ == "__main__":
    update_template_var("template.docx")
