import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_alignment(path):
    print(f"Processing {path}...")
    doc = docx.Document(path)
    count = 0
    
    # Target placeholders
    targets = [
        '{{ v.description }}', 
        '{{ v.impact }}', 
        '{{ v.recommendation }}', 
        '{{ v.affected_url }}'
    ]
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    # Check if paragraph contains any of our targets
                    if any(t in p.text for t in targets):
                        # Force Left Alignment
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        count += 1
                        print(f"Fixed alignment for paragraph containing: {p.text[:30]}...")
                        
    doc.save(path)
    print(f"Done. Fixed {count} locations.")

if __name__ == "__main__":
    fix_alignment("template.docx")
