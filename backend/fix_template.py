import docx
import os

def delete_row(row):
    """Safely remove a row from a table."""
    tbl = row._element.getparent()
    tbl.remove(row._element)

def fix_template(path):
    print(f"Processing {path}...")
    doc = docx.Document(path)
    
    rows_to_delete = []
    
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            text = "".join(cell.text for cell in row.cells).strip()
            
            # Check for generic start loop tag
            if "{%" in text and "for " in text and "%}" in text and "endfor" not in text:
                # This is a Loop Start Row
                print(f"Found Loop Start at Row {i}: {text}")
                
                # Move tag to the NEXT row (the data row), first cell
                if i + 1 < len(table.rows):
                    data_row = table.rows[i+1]
                    # We prepend the tag to the first paragraph of the first cell
                    # Using insert_paragraph_before to avoid messing up run formatting of existing text
                    first_cell = data_row.cells[0]
                    p = first_cell.paragraphs[0]
                    p.insert_paragraph_before(text)
                    
                    rows_to_delete.append(row)
            
            # Check for generic end loop tag
            elif "{%" in text and "endfor" in text and "%}" in text:
                # This is a Loop End Row
                print(f"Found Loop End at Row {i}: {text}")
                
                # Move tag to the PREVIOUS row (the data row), last cell
                if i - 1 >= 0:
                    data_row = table.rows[i-1]
                    last_cell = data_row.cells[-1]
                    last_cell.add_paragraph(text)
                    
                    rows_to_delete.append(row)

    # Delete rows in reverse order to avoid index shifting (though strict object ref usually works, reverse is safer)
    for row in reversed(rows_to_delete):
        try:
            delete_row(row)
            print("Deleted a control row.")
        except Exception as e:
            print(f"Error deleting row: {e}")

    doc.save(path)
    print("Template saved.")

if __name__ == "__main__":
    fix_template("template.docx")
