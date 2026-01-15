import docx

def repair(path):
    print(f"Repairing {path}...")
    doc = docx.Document(path)
    
    # ---------------------------------------------------------
    # Table 1: Summary Table
    # Row 1 contains {{ v.summary_index }}
    # We need to wrap this row in the loop.
    # ---------------------------------------------------------
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        if len(t1.rows) > 1:
            row = t1.rows[1] # Data row
            
            # Cell 0: Prepend {% for v in vulns %}
            # We insert it into the first paragraph, first run if possible, or prepend text
            c0 = row.cells[0]
            if c0.paragraphs:
                p = c0.paragraphs[0]
                # Prepend text to the first run to keep it inline
                if p.runs:
                    p.runs[0].text = "{% for v in vulns %} " + p.runs[0].text
                else:
                    p.add_run("{% for v in vulns %} ")
            
            # Cell Last: Append {% endfor %}
            cLast = row.cells[-1]
            if cLast.paragraphs:
                p = cLast.paragraphs[-1] # Last paragraph
                p.add_run(" {% endfor %}")
            
            print("Repaired Table 1 Row 1")

    # ---------------------------------------------------------
    # Table 2: Detail Table (Assumed)
    # Row 0 contains {{ v.detail_index }}
    # ---------------------------------------------------------
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        if len(t2.rows) > 0:
            row = t2.rows[0] # Data row
            
            # Cell 0: Prepend {% for v in vulns %}
            c0 = row.cells[0]
            if c0.paragraphs:
                p = c0.paragraphs[0]
                if p.runs:
                    p.runs[0].text = "{% for v in vulns %} " + p.runs[0].text
                else:
                    p.add_run("{% for v in vulns %} ")

            # Cell Last: Append {% endfor %}
            # Table 2 might be a header or single row table.
            cLast = row.cells[-1]
            if cLast.paragraphs:
                p = cLast.paragraphs[-1]
                p.add_run(" {% endfor %}")
                
            print("Repaired Table 2 Row 0")

    doc.save(path)
    print("Template repaired.")

if __name__ == "__main__":
    repair("template.docx")
