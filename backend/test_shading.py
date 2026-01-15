from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

def test_shading():
    print("Testing shading logic...")
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    
    # Header
    table.rows[0].cells[0].text = "S/N"
    table.rows[0].cells[1].text = "Vulnerability"
    table.rows[0].cells[2].text = "Severity"
    
    # Data
    table.rows[1].cells[2].text = "High"
    
    def set_cell_shading(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        existing_shd = tcPr.find(qn('w:shd'))
        if existing_shd is not None:
            tcPr.remove(existing_shd)
        
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    # Logic from main.py
    column_index_severity = -1
    for table in doc.tables:
        if len(table.rows) > 0:
            header_cells = table.rows[0].cells
            found_severity = False
            for idx, cell in enumerate(header_cells):
                if "Severity" in cell.text:
                    column_index_severity = idx
                    found_severity = True
                    break
            
            if found_severity:
                for row in table.rows[1:]:
                     cell = row.cells[column_index_severity]
                     print(f"Shading cell with text: {cell.text}")
                     set_cell_shading(cell, "FF0000")
                     
    print("Shading applied successfully.")

if __name__ == "__main__":
    try:
        test_shading()
    except Exception as e:
        print(f"FAILED: {e}")
        import traceback
        traceback.print_exc()
