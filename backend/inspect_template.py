import docx

def inspect(path):
    doc = docx.Document(path)
    for t_idx, table in enumerate(doc.tables):
        print(f"--- Table {t_idx} ---")
        for r_idx, row in enumerate(table.rows):
            cell_texts = []
            for cell in row.cells:
                # Print paragraphs to see exact placement
                p_texts = [p.text for p in cell.paragraphs]
                cell_texts.append(" | ".join(p_texts))
            print(f"Row {r_idx}: {cell_texts}")

if __name__ == "__main__":
    inspect("template.docx")
