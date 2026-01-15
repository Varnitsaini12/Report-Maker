import docx
import sys

def verify(path):
    print(f"Inspecting {path}...")
    doc = docx.Document(path)
    for t_idx, table in enumerate(doc.tables):
        print(f"--- Table {t_idx} ---")
        for r_idx, row in enumerate(table.rows):
            try:
                if not row.cells:
                    print(f"R{r_idx}: NO CELLS")
                    continue
                
                c0_paras = [p.text for p in row.cells[0].paragraphs]
                cLast_paras = [p.text for p in row.cells[-1].paragraphs]
                
                # Check for tags
                full_text = " ".join(c0_paras + cLast_paras)
                if "{%" in full_text:
                    print(f"*** FOUND TAGS IN T{t_idx} R{r_idx} ***")
                    print(f"  C0: {c0_paras}")
                    print(f"  CL: {cLast_paras}")
                else:
                    # Print normal rows sparingly or if interesting
                    if r_idx < 2: # Print headers
                        print(f"R{r_idx} C0: {c0_paras}")
            except Exception as e:
                print(f"R{r_idx}: Error {e}")

if __name__ == "__main__":
    verify("template.docx")
